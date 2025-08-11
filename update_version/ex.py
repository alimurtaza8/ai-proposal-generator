

# main.py - Fixed Enhanced FastAPI Backend for RFP Proposal Generator
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any, Tuple, Union
import os
import uuid
import shutil
from datetime import datetime
import json
import asyncio
from dotenv import load_dotenv
load_dotenv()
import requests
import base64
from urllib.parse import quote

# Document processing imports
import PyPDF2
import docx
from PIL import Image
import pytesseract
import re

# AI integration imports
import google.generativeai as genai

# Document generation imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import traceback
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import time

# Arabic text processing
try:
    import arabic_reshaper
    from bidi.algorithm import get_display
    ARABIC_SUPPORT = True
except ImportError:
    ARABIC_SUPPORT = False
    print("WARNING: Arabic text processing libraries not installed. Arabic support disabled.")

# PDF font registration
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# --- Project Root Definition ---
# To handle file paths correctly regardless of where the script is run
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

def resolve_project_path(relative_path):
    """Resolves a path from the project root."""
    if not relative_path:
        return None
    return os.path.join(PROJECT_ROOT, relative_path)

# Create FastAPI app
app = FastAPI(title="RFP Proposal Generator API", version="2.1.0")

# New endpoint to serve logos from the root of the backend directory
@app.get("/logos/{logo_name}")
async def get_logo(logo_name: str):
    # Security: only allow specific, known logo files to be served
    allowed_logos = ["LOGO-04.png", "LOGO-05.png"]
    if logo_name not in allowed_logos:
        raise HTTPException(status_code=404, detail="Logo not found")
    
    file_path = os.path.join(os.path.dirname(__file__), logo_name)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found on server")
    
    return FileResponse(file_path)

# --- Font Registration ---
AMIRI_FONT_PATH = None
# Check for common variations of the font file name
potential_paths = [
    'backend/fonts/Amiri-Regular.ttf',
    'backend/fonts/amiri-Reggular.ttf',
    'backend/fonts/Amiri-regular.ttf',
    'backend/Amiri-Regular.ttf',
    'fonts/Amiri-Regular.ttf',
    './Amiri-Regular.ttf'
]

for path in potential_paths:
    if os.path.exists(path):
        AMIRI_FONT_PATH = path
        break

if AMIRI_FONT_PATH and ARABIC_SUPPORT:
    try:
        pdfmetrics.registerFont(TTFont('Amiri', AMIRI_FONT_PATH))
        pdfmetrics.registerFontFamily('AmiriFamily', normal='Amiri', bold='Amiri', italic='Amiri', boldItalic='Amiri')
        print(f"Amiri font registered successfully from {AMIRI_FONT_PATH}.")
    except Exception as e:
        print(f"FATAL: Could not register Amiri font from {AMIRI_FONT_PATH}: {e}")
        AMIRI_FONT_PATH = None
else:
    if not ARABIC_SUPPORT:
        print("WARNING: Arabic text processing libraries not available. Arabic PDF generation will be disabled.")
    else:
        print("WARNING: Amiri font file not found. Arabic PDF generation will be disabled.")

# Translations for static text
TRANSLATIONS = {
    'en': {
        'technical_proposal': 'Technical Proposal',
        'prepared_for': 'Prepared for',
        'table_of_contents': 'Table of Contents',
        'page': 'Page',
        'of': 'of'
    },
    'ar': {
        'technical_proposal': 'مقترح فني',
        'prepared_for': 'مُعد لـ',
        'table_of_contents': 'جدول المحتويات',
        'page': 'صفحة',
        'of': 'من'
    }
}

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configuration
UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

# Ensure directories exist
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Enhanced Data models
class ProposalRequest(BaseModel):
    proposal_type: str
    sector: str
    company_name: str
    selected_sections: Optional[List[str]] = None
    phases_to_include: Optional[List[int]] = None
    output_format: Optional[str] = "all"
    logo_top_left_path: Optional[str] = None
    logo_bottom_right_path: Optional[str] = None
    language: Optional[str] = "en"
    # Enhanced document analysis fields
    special_document_content: Optional[str] = None
    additional_documents_content: Optional[List[str]] = None
    special_document_insights: Optional[str] = None
    additional_documents_insights: Optional[str] = None

class ProposalResponse(BaseModel):
    job_id: str
    status: str
    message: str
    files: Optional[List[str]] = None
    progress: Optional[int] = None
    structure_info: Optional[Dict[str, Any]] = None

class ExtractedStructure(BaseModel):
    sections: List[Dict[str, Any]]
    requirements: List[str]
    scope: str
    timeline: Optional[str] = None

# Mermaid Diagram Generator for Professional Visuals
class MermaidDiagramGenerator:
    def __init__(self):
        self.base_url = "https://mermaid.ink"
    
    def generate_timeline_mermaid(self, project_name: str, phases: list) -> str:
        """Generate Mermaid syntax for timeline diagram"""
        mermaid_code = f"""gantt
    title {project_name} Timeline
    dateFormat YYYY-MM-DD
    
"""
        for phase in phases:
            mermaid_code += f"    section {phase['name']}\n"
            mermaid_code += f"    {phase['task']} :a{phase.get('id', 1)}, {phase['start']}, {phase['duration']}\n"
        
        # Generate viewable URL
        encoded = base64.b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
        url = f"https://mermaid.live/edit#pako:{encoded}"
        
        return f"""
**Project Timeline Visualization:**

```mermaid
{mermaid_code}
```

[View Interactive Timeline]({url})

The timeline above shows the complete project schedule with dependencies and milestones clearly marked.
"""
    
    def generate_architecture_mermaid(self, system_name: str, components: list) -> str:
        """Generate Mermaid syntax for architecture diagram"""
        mermaid_code = f"""graph TD
    %% {system_name} Architecture
    
"""
        for component in components:
            mermaid_code += f"    {component['id']}[{component['name']}]\n"
        
        mermaid_code += "\n    %% Connections\n"
        for component in components:
            if 'connects_to' in component:
                for connection in component['connects_to']:
                    mermaid_code += f"    {component['id']} --> {connection}\n"
        
        encoded = base64.b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
        url = f"https://mermaid.live/edit#pako:{encoded}"
        
        return f"""
**System Architecture Diagram:**

```mermaid
{mermaid_code}
```

[View Interactive Architecture]({url})

This diagram illustrates the complete system architecture with all components and their relationships.
"""

    def generate_modular_mermaid(self, solution_name: str, modules: list) -> str:
        """Generate Mermaid syntax for modular design"""
        mermaid_code = f"""graph LR
    %% {solution_name} Modular Design
    
"""
        for module in modules:
            mermaid_code += f"    {module['id']}[{module['name']}]\n"
            if 'submodules' in module:
                for sub in module['submodules']:
                    mermaid_code += f"    {module['id']} --> {sub['id']}[{sub['name']}]\n"
        
        encoded = base64.b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
        url = f"https://mermaid.live/edit#pako:{encoded}"
        
        return f"""
**Modular Solution Design:**

```mermaid
{mermaid_code}
```

[View Interactive Modular Design]({url})

This diagram shows how our modular solution architecture enables scalability and maintainability.
"""

# Enhanced Section Structure for Dynamic Generation
class Section:
    def __init__(self, key: str, title: str, level: int = 1, subsections: List['Section'] = None, content_requirements: List[str] = None):
        self.key = key
        self.title = title
        self.level = level
        self.subsections = subsections or []
        self.content_requirements = content_requirements or []
        self.number = ""
        self.is_dynamic = False
    
    def add_subsection(self, section: 'Section'):
        self.subsections.append(section)
        return self
    
    def to_dict(self):
        return {
            "key": self.key,
            "title": self.title,
            "level": self.level,
            "number": self.number,
            "is_dynamic": self.is_dynamic,
            "content_requirements": self.content_requirements,
            "subsections": [sub.to_dict() for sub in self.subsections]
        }

# Visualization Document Generator for Mermaid Diagrams
class VisualizationDocumentGenerator:
    def __init__(self):
        self.diagram_generator = MermaidDiagramGenerator()
    
    def generate_visualization_html(self, content: dict, structure: List[Section], company_name: str, job_id: str, request: 'ProposalRequest') -> str:
        """Generate a separate HTML file with all visualizations"""
        
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{company_name} - Project Visualizations</title>
    <script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f8f9fa;
        }}
        .header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 2rem;
            border-radius: 10px;
            margin-bottom: 2rem;
            text-align: center;
        }}
        .visualization-section {{
            background: white;
            padding: 2rem;
            margin-bottom: 2rem;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        .mermaid {{
            text-align: center;
            margin: 2rem 0;
        }}
        h1 {{ margin: 0; font-size: 2.5rem; }}
        h2 {{ color: #667eea; border-bottom: 2px solid #667eea; padding-bottom: 0.5rem; }}
        h3 {{ color: #555; }}
        .description {{
            background: #e3f2fd;
            padding: 1rem;
            border-radius: 5px;
            margin: 1rem 0;
            border-left: 4px solid #2196F3;
        }}
        .timestamp {{
            text-align: center;
            color: #666;
            font-size: 0.9rem;
            margin-top: 2rem;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Project Visualizations</h1>
        <p>Interactive Diagrams and Timeline for {company_name}</p>
        <p>Proposal Type: {request.proposal_type.title()} | Sector: {request.sector.title()}</p>
    </div>
"""
        
        # Generate visualizations for each relevant section
        for section in structure:
            if self._should_include_visualization(section):
                html_content += self._generate_section_visualization(section, content, request)
        
        html_content += f"""
    <div class="timestamp">
        Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
    </div>
    
    <script>
        mermaid.initialize({{ startOnLoad: true, theme: 'default' }});
    </script>
</body>
</html>
"""
        
        # Save HTML file
        visualization_filename = f"visualization_{job_id}_{company_name.replace(' ', '_')}.html"
        visualization_path = os.path.join(OUTPUT_DIR, visualization_filename)
        
        with open(visualization_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return visualization_filename
    
    def _should_include_visualization(self, section: Section) -> bool:
        """Check if section should have visualizations"""
        visual_keywords = ['timeline', 'architecture', 'modular', 'implementation', 'deliverables', 'structure']
        return any(keyword in section.key.lower() for keyword in visual_keywords)
    
    def _generate_section_visualization(self, section: Section, content: dict, request: 'ProposalRequest') -> str:
        """Generate visualization for a specific section"""
        section_html = f"""
    <div class="visualization-section">
        <h2>{section.number}. {section.title}</h2>
"""
        
        # Generate appropriate diagram based on section type
        if 'timeline' in section.key.lower():
            section_html += self._generate_timeline_visualization(section, content, request)
        elif 'architecture' in section.key.lower():
            section_html += self._generate_architecture_visualization(section, content, request)
        elif 'modular' in section.key.lower():
            section_html += self._generate_modular_visualization(section, content, request)
        elif 'implementation' in section.key.lower():
            section_html += self._generate_implementation_visualization(section, content, request)
        elif 'deliverables' in section.key.lower():
            section_html += self._generate_deliverables_visualization(section, content, request)
        
        section_html += "</div>\n"
        return section_html
    
    def _generate_timeline_visualization(self, section: Section, content: dict, request: 'ProposalRequest') -> str:
        """Generate timeline visualization"""
        return f"""
        <h3>Project Timeline</h3>
        <div class="description">
            This interactive timeline shows the complete project schedule with all phases, milestones, and dependencies.
        </div>
        <div class="mermaid">
gantt
    title {request.company_name} Project Timeline
    dateFormat YYYY-MM-DD
    
    section Planning Phase
    Requirements Analysis    :active, req, 2024-01-01, 2w
    System Design          :design, after req, 2w
    
    section Development Phase
    Core Development       :dev, after design, 6w
    Integration Testing    :test, after dev, 2w
    
    section Deployment Phase
    User Training          :training, after test, 1w
    Go-Live               :golive, after training, 1w
        </div>
"""
    
    def _generate_architecture_visualization(self, section: Section, content: dict, request: 'ProposalRequest') -> str:
        """Generate architecture visualization"""
        return f"""
        <h3>System Architecture</h3>
        <div class="description">
            This diagram illustrates the complete system architecture with all components and their relationships.
        </div>
        <div class="mermaid">
graph TD
    A[User Interface Layer] --> B[API Gateway]
    B --> C[Authentication Service]
    B --> D[Business Logic Layer]
    D --> E[Data Access Layer]
    E --> F[(Primary Database)]
    E --> G[(Cache Layer)]
    D --> H[External Services]
    I[Load Balancer] --> A
    J[CDN] --> I
        </div>
"""
    
    def _generate_modular_visualization(self, section: Section, content: dict, request: 'ProposalRequest') -> str:
        """Generate modular design visualization"""
        return f"""
        <h3>Modular Solution Design</h3>
        <div class="description">
            This diagram shows how our modular solution architecture enables scalability and maintainability.
        </div>
        <div class="mermaid">
graph LR
    Core[Core System] --> UserMgmt[User Management Module]
    Core --> DataProc[Data Processing Module]
    Core --> Reports[Reporting Module]
    Core --> Integration[Integration Module]
    
    UserMgmt --> Auth[Authentication]
    UserMgmt --> Profile[Profile Management]
    UserMgmt --> Permissions[Permissions]
    
    DataProc --> Ingestion[Data Ingestion]
    DataProc --> Validation[Data Validation]
    DataProc --> Transform[Data Transformation]
    
    Reports --> Dashboard[Dashboard]
    Reports --> Analytics[Analytics Engine]
    Reports --> Export[Export Functions]
        </div>
"""
    
    def _generate_implementation_visualization(self, section: Section, content: dict, request: 'ProposalRequest') -> str:
        """Generate implementation process visualization"""
        return f"""
        <h3>Implementation Process Flow</h3>
        <div class="description">
            This flowchart shows the step-by-step implementation process with decision points and deliverables.
        </div>
        <div class="mermaid">
flowchart TD
    A[Project Kickoff] --> B{{Requirements Review}}
    B --> C[System Design]
    C --> D[Development Sprint 1]
    D --> E[Testing & QA]
    E --> F{{Quality Gate}}
    F -->|Pass| G[Development Sprint 2]
    F -->|Fail| D
    G --> H[Integration Testing]
    H --> I[User Acceptance Testing]
    I --> J{{UAT Approval}}
    J -->|Approved| K[Production Deployment]
    J -->|Changes Required| G
    K --> L[Go-Live Support]
        </div>
"""
    
    def _generate_deliverables_visualization(self, section: Section, content: dict, request: 'ProposalRequest') -> str:
        """Generate deliverables breakdown visualization"""
        return f"""
        <h3>Project Deliverables Structure</h3>
        <div class="description">
            This diagram shows all project deliverables organized by phase and their interdependencies.
        </div>
        <div class="mermaid">
graph TD
    subgraph "Phase 1: Planning"
        A[Requirements Document]
        B[System Architecture]
        C[Project Plan]
    end
    
    subgraph "Phase 2: Development"
        D[Core System]
        E[User Interface]
        F[API Documentation]
    end
    
    subgraph "Phase 3: Testing"
        G[Test Results]
        H[User Manual]
        I[Training Materials]
    end
    
    subgraph "Phase 4: Deployment"
        J[Production System]
        K[Support Documentation]
        L[Maintenance Plan]
    end
    
    A --> D
    B --> D
    B --> E
    C --> D
    D --> G
    E --> G
    G --> J
    H --> I
    I --> L
        </div>
"""

# AI Configuration
class AIConfig:
    def __init__(self):
        self.gemini_api_key = os.getenv("GEMINI_API_KEY")
        
        if self.gemini_api_key:
            try:
                genai.configure(api_key=self.gemini_api_key)
                self.model = genai.GenerativeModel('gemini-1.5-flash')
                print("Gemini API configured successfully!")
            except Exception as e:
                print(f"Failed to configure Gemini API: {e}")
                self.model = None
        else:
            print("GEMINI_API_KEY not found in environment variables")
            self.model = None

ai_config = AIConfig()

# Enhanced Document Processor with Intelligent Structure Extraction
class EnhancedDocumentProcessor:
    def __init__(self):
        self.section_patterns = [
            r'^(\d+\.?\d*\.?\d*)\s+(.+)',
            r'^([A-Z][A-Z\s]+)$',
            r'^([A-Z]\.\d+)\s+(.+)',
            r'^(Chapter\s+\d+)[\s\-:]+(.+)',
            r'^(Section\s+\d+)[\s\-:]+(.+)',
            r'^(Part\s+[IVX]+)[\s\-:]+(.+)',
            r'^(Appendix\s+[A-Z])[\s\-:]+(.+)',
        ]
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += f"\n--- PAGE {page_num + 1} ---\n{page_text}\n"
                return text
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                style_info = f"[STYLE:{paragraph.style.name}]" if paragraph.style.name != 'Normal' else ""
                text += f"{style_info}{paragraph.text}\n"
            return text
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading Word document: {str(e)}")
    
    def extract_text_from_image(self, file_path: str) -> str:
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            return f"Image content extracted (OCR processing error: {str(e)})"
    
    def process_file(self, file_path: str) -> tuple:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            text = self.extract_text_from_pdf(file_path)
            structure = self.extract_structure_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            text = self.extract_text_from_docx(file_path)
            structure = self.extract_structure_from_docx(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png']:
            text = self.extract_text_from_image(file_path)
            structure = self.extract_structure_from_text(text)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_ext}")
        
        return text, structure

    def extract_structure_from_pdf(self, file_path: str) -> ExtractedStructure:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""
                for page in pdf_reader.pages:
                    full_text += page.extract_text() + "\n"
                
                return self.analyze_document_structure(full_text)
        except Exception as e:
            print(f"Error extracting PDF structure: {e}")
            return ExtractedStructure(sections=[], requirements=[], scope="")

    def extract_structure_from_docx(self, file_path: str) -> ExtractedStructure:
        try:
            doc = docx.Document(file_path)
            sections = []
            requirements = []
            scope_parts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                style_name = paragraph.style.name
                
                if style_name.startswith('Heading'):
                    level = int(style_name.replace('Heading ', ''))
                    sections.append({
                        'title': text,
                        'level': level,
                        'key': self.create_key_from_title(text),
                        'content_type': 'heading'
                    })
                
                if any(req_word in text.lower() for req_word in ['must', 'shall', 'requirement', 'mandatory', 'essential']):
                    requirements.append(text)
                
                if any(scope_word in text.lower() for scope_word in ['scope', 'objective', 'purpose', 'goal']):
                    scope_parts.append(text)
            
            return ExtractedStructure(
                sections=sections,
                requirements=requirements,
                scope=' '.join(scope_parts[:3])
            )
        except Exception as e:
            print(f"Error extracting DOCX structure: {e}")
            return ExtractedStructure(sections=[], requirements=[], scope="")

    def extract_structure_from_text(self, text: str) -> ExtractedStructure:
        return self.analyze_document_structure(text)
    
    def analyze_document_structure(self, text: str) -> ExtractedStructure:
        lines = text.split('\n')
        sections = []
        requirements = []
        scope_parts = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            section_info = self.identify_section_heading(line)
            if section_info:
                sections.append(section_info)
            
            if any(req_word in line.lower() for req_word in 
                   ['must', 'shall', 'requirement', 'mandatory', 'essential', 'required']):
                requirements.append(line)
            
            if any(scope_word in line.lower() for scope_word in 
                   ['scope', 'objective', 'purpose', 'goal', 'deliverable']):
                scope_parts.append(line)
        
        return ExtractedStructure(
            sections=sections,
            requirements=requirements,
            scope=' '.join(scope_parts[:5])
        )
    
    def identify_section_heading(self, line: str) -> Optional[Dict[str, Any]]:
        for pattern in self.section_patterns:
            match = re.match(pattern, line)
            if match:
                if len(match.groups()) == 2:
                    number, title = match.groups()
                    level = self.determine_level_from_number(number)
                else:
                    number = ""
                    title = match.group(1)
                    level = 1
                
                return {
                    'title': title.strip(),
                    'level': level,
                    'key': self.create_key_from_title(title.strip()),
                    'number': number.strip(),
                    'content_type': 'heading'
                }
        
        if (line.isupper() and len(line) > 5 and len(line) < 100 and 
            not any(char.isdigit() for char in line[:10])):
            return {
                'title': line,
                'level': 1,
                'key': self.create_key_from_title(line),
                'number': "",
                'content_type': 'heading'
            }
        
        return None
    
    def determine_level_from_number(self, number: str) -> int:
        if not number:
            return 1
        
        dot_count = number.count('.')
        if dot_count > 0:
            return min(dot_count + 1, 3)
        
        if number.isdigit():
            return 1
        
        return 1
    
    def create_key_from_title(self, title: str) -> str:
        key = re.sub(r'[^\w\s]', '', title.lower())
        key = re.sub(r'\s+', '_', key)
        key = key.strip('_')
        return key[:50]

# Enhanced AI Content Generator with Dynamic Structure
class EnhancedAIContentGenerator:
    def __init__(self):
        self.processor = EnhancedDocumentProcessor()
    
    async def analyze_rfp_and_generate_structure(self, rfp_text: str, extracted_structure: ExtractedStructure, request: ProposalRequest) -> List[Section]:
        try:
            if ai_config.model:
                structure = await self._generate_structure_with_ai(rfp_text, extracted_structure, request)
            else:
                structure = self._generate_fallback_structure(extracted_structure)
            
            self.number_sections(structure)
            return structure
            
        except Exception as e:
            print(f"Error generating structure: {e}")
            return self._generate_fallback_structure(extracted_structure)
    
    async def _generate_structure_with_ai(self, rfp_text: str, extracted_structure: ExtractedStructure, request: ProposalRequest) -> List[Section]:
        prompt = f"""
Analyze this RFP document and generate an appropriate proposal structure.

RFP CONTENT:
{rfp_text[:3000]}

EXTRACTED SECTIONS FROM RFP:
{json.dumps([s for s in extracted_structure.sections[:10]], indent=2)}

EXTRACTED REQUIREMENTS:
{json.dumps(extracted_structure.requirements[:5], indent=2)}

SCOPE:
{extracted_structure.scope}

Generate a comprehensive proposal structure that addresses all RFP requirements. 
Create sections and subsections that logically respond to the RFP's needs.
The titles for the sections and subsections MUST be in {request.language}. The keys should remain in English.

Respond with ONLY a valid JSON array of sections with this structure:
[
  {{
    "key": "executive_summary",
    "title": "Executive Summary",
    "level": 1,
    "content_requirements": ["Brief overview", "Key benefits", "Recommendations"]
  }},
  {{
    "key": "understanding_requirements", 
    "title": "Understanding of Requirements",
    "level": 1,
    "content_requirements": ["RFP analysis", "Key challenges identified"],
    "subsections": [
      {{
        "key": "technical_requirements",
        "title": "Technical Requirements Analysis", 
        "level": 2,
        "content_requirements": ["Technical specifications", "Compliance requirements"]
      }}
    ]
  }}
]

Include 8-12 main sections with relevant subsections. Focus on BUSINESS VALUE FIRST, then technical details:
1. Executive Summary
2. Understanding of Requirements  
3. Deliverables and Expected Outcomes
4. Pricing and Investment Structure 
5. Proposed Solution/Approach
6. Technical Specifications
7. Implementation Plan and Timeline
8. Team and Qualifications
9. Risk Management and Mitigation
10. Quality Assurance and Success Metrics
11. Support and Maintenance
12. Conclusion and Next Steps

CRITICAL: Prioritize deliverables and pricing (sections 3-4) BEFORE technical approach. Clients need to understand WHAT they get and HOW MUCH it costs before diving into HOW it will be built.

Ensure each section has appropriate content_requirements that guide content generation.
"""
        
        try:
            response = ai_config.model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.7,
                    top_k=40,
                    top_p=0.95,
                    max_output_tokens=4000,
                )
            )
            
            if response.text:
                cleaned_text = response.text.strip()
                if cleaned_text.startswith('```json'):
                    cleaned_text = cleaned_text.replace('```json', '').replace('```', '').strip()
                
                structure_data = json.loads(cleaned_text)
                return self._convert_json_to_sections(structure_data)
        
        except Exception as e:
            print(f"AI structure generation failed: {e}")
        
        return self._generate_fallback_structure(extracted_structure)
    
    def _convert_json_to_sections(self, structure_data: List[Dict]) -> List[Section]:
        sections = []
        
        for item in structure_data:
            section = Section(
                key=item.get('key', ''),
                title=item.get('title', ''),
                level=item.get('level', 1),
                content_requirements=item.get('content_requirements', [])
            )
            section.is_dynamic = True
            
            if 'subsections' in item:
                for sub_item in item['subsections']:
                    subsection = Section(
                        key=sub_item.get('key', ''),
                        title=sub_item.get('title', ''),
                        level=sub_item.get('level', 2),
                        content_requirements=sub_item.get('content_requirements', [])
                    )
                    subsection.is_dynamic = True
                    section.add_subsection(subsection)
            
            sections.append(section)
        
        return sections
    
    def _generate_fallback_structure(self, extracted_structure: ExtractedStructure) -> List[Section]:
        sections = []
        
        # 1. Executive Summary
        sections.append(Section("executive_summary", "Executive Summary", 1, 
                              content_requirements=["Project overview", "Key benefits", "Value proposition", "Investment summary"]))
        
        # 2. Understanding of Requirements
        sections.append(Section("understanding_requirements", "Understanding of Requirements", 1,
                              content_requirements=["RFP analysis", "Key challenges", "Scope clarification", "Success criteria"]))
        
        # 3. Deliverables and Expected Outcomes (PRIORITY: Before technical details)
        deliverables_section = Section("deliverables_outcomes", "Deliverables and Expected Outcomes", 1,
                                     content_requirements=["Key deliverables", "Expected outcomes", "Success metrics", "Value to client"])
        deliverables_section.add_subsection(Section("primary_deliverables", "Primary Deliverables", 2,
                                                   content_requirements=["Main outputs", "Quality standards", "Acceptance criteria"]))
        deliverables_section.add_subsection(Section("expected_outcomes", "Expected Business Outcomes", 2,
                                                   content_requirements=["Business impact", "Performance improvements", "ROI expectations"]))
        sections.append(deliverables_section)
        
        # 4. Pricing and Investment Structure (PRIORITY: Before technical details)
        pricing_section = Section("pricing_investment", "Pricing and Investment Structure", 1,
                                content_requirements=["Investment overview", "Pricing model", "Value justification", "Payment terms"])
        pricing_section.add_subsection(Section("investment_summary", "Investment Summary", 2,
                                              content_requirements=["Total investment", "Cost breakdown", "Value proposition"]))
        pricing_section.add_subsection(Section("pricing_model", "Pricing Model", 2,
                                              content_requirements=["Pricing structure", "Payment schedule", "Terms and conditions"]))
        sections.append(pricing_section)
        
        # 5. Proposed Solution/Approach (NOW after business value is clear)
        solution_section = Section("proposed_solution", "Proposed Solution and Approach", 1,
                                 content_requirements=["Solution overview", "Approach methodology", "Innovation highlights"])
        solution_section.add_subsection(Section("solution_overview", "Solution Overview", 2,
                                               content_requirements=["High-level approach", "Key features", "Differentiators"]))
        solution_section.add_subsection(Section("methodology", "Implementation Methodology", 2,
                                               content_requirements=["Process framework", "Best practices", "Quality approach"]))
        sections.append(solution_section)
        
        # 6. Technical Specifications (After business value is established)
        technical_section = Section("technical_specifications", "Technical Specifications", 1,
                                  content_requirements=["Technical requirements", "Architecture overview", "Technology stack"])
        technical_section.add_subsection(Section("technical_approach", "Technical Approach", 2,
                                                content_requirements=["Architecture", "Technologies", "Standards", "Integration"]))
        technical_section.add_subsection(Section("technical_requirements", "Technical Requirements", 2,
                                                content_requirements=["System requirements", "Performance specifications", "Compliance standards"]))
        sections.append(technical_section)
        
        # 7. Implementation Plan and Timeline
        implementation_section = Section("implementation_plan", "Implementation Plan and Timeline", 1,
                                        content_requirements=["Project phases", "Timeline", "Milestones", "Resource allocation"])
        implementation_section.add_subsection(Section("project_phases", "Project Phases", 2,
                                                     content_requirements=["Phase breakdown", "Phase deliverables", "Dependencies"]))
        implementation_section.add_subsection(Section("timeline_milestones", "Timeline and Milestones", 2,
                                                     content_requirements=["Project schedule", "Key dates", "Critical path", "Visual timeline diagram"]))
        implementation_section.add_subsection(Section("project_timeline_visual", "Project Timeline Visualization", 2,
                                                     content_requirements=["Gantt-style timeline", "Phase dependencies", "Milestone markers", "Resource allocation visual"]))
        sections.append(implementation_section)
        
        # 8. Team and Qualifications
        sections.append(Section("team_qualifications", "Team and Qualifications", 1,
                              content_requirements=["Team structure", "Key personnel", "Relevant experience", "Certifications"]))
        
        # 9. Risk Management and Mitigation
        sections.append(Section("risk_management", "Risk Management and Mitigation", 1,
                              content_requirements=["Risk identification", "Mitigation strategies", "Contingency plans", "Risk monitoring"]))
        
        # 10. Quality Assurance and Success Metrics
        sections.append(Section("quality_assurance", "Quality Assurance and Success Metrics", 1,
                              content_requirements=["QA processes", "Testing procedures", "Success metrics", "Performance monitoring"]))
        
        # 11. Support and Maintenance
        sections.append(Section("support_maintenance", "Support and Maintenance", 1,
                              content_requirements=["Support model", "Maintenance approach", "SLA commitments", "Ongoing services"]))
        
        # 12. Conclusion and Next Steps
        sections.append(Section("conclusion", "Conclusion and Next Steps", 1,
                              content_requirements=["Summary", "Next steps", "Call to action", "Contact information"]))
        
        return sections
    
    def number_sections(self, sections: List[Section], parent_number: str = ""):
        for i, section in enumerate(sections):
            if parent_number:
                section.number = f"{parent_number}.{i+1}"
            else:
                section.number = str(i+1)
            
            if section.subsections:
                self.number_sections(section.subsections, section.number)
    
    def flatten_sections(self, sections: List[Section], flat_list: List[Section] = None) -> List[Section]:
        if flat_list is None:
            flat_list = []
        
        for section in sections:
            flat_list.append(section)
            if section.subsections:
                self.flatten_sections(section.subsections, flat_list)
        
        return flat_list
    
    async def generate_proposal_content(self, rfp_text: str, structure: List[Section], request: ProposalRequest) -> dict:
        try:
            flat_sections = self.flatten_sections(structure)
            sections_to_generate = self._filter_sections(flat_sections, request)
            
            if ai_config.model:
                content = await self._generate_content_with_ai(rfp_text, sections_to_generate, request)
            else:
                content = self._generate_mock_content(sections_to_generate, request)
            
            return content
        except Exception as e:
            print(f"Error generating content: {e}")
            raise HTTPException(status_code=500, detail=f"Error generating content: {str(e)}")
    
    def _filter_sections(self, sections: List[Section], request: ProposalRequest) -> List[Section]:
        if not request.selected_sections:
            return sections
        
        return [s for s in sections if s.key in request.selected_sections]
    
    async def _generate_content_with_ai(self, rfp_text: str, sections: List[Section], request: ProposalRequest) -> dict:
        content = {}
        tasks = []

        for section in sections:
            tasks.append(self._generate_single_section_content(rfp_text, section, request))

        results = await asyncio.gather(*tasks)

        for section, section_content in zip(sections, results):
            if section_content:
                content[section.key] = section_content

        return content

    async def _generate_single_section_content(self, rfp_text: str, section: Section, request: ProposalRequest) -> Optional[str]:
        section_info = f"{section.number}. {section.title} (Level {section.level})"
        if section.content_requirements:
            section_info += f" - Requirements: {', '.join(section.content_requirements)}"

        # Include insights from special and additional documents
        insights_context = ""
        if request.special_document_insights:
            insights_context += f"\n\nSPECIAL DOCUMENT INSIGHTS:\n{request.special_document_insights[:1000]}\n"
        
        if request.additional_documents_insights:
            insights_context += f"\n\nADDITIONAL DOCUMENTS INSIGHTS:\n{request.additional_documents_insights[:1000]}\n"

        # Add specific guidance based on section type
        section_guidance = ""
        if "deliverables" in section.key.lower() or "outcomes" in section.key.lower():
            section_guidance = """
SPECIFIC GUIDANCE FOR DELIVERABLES/OUTCOMES SECTION:
- Focus on CONCRETE, MEASURABLE deliverables that the client will receive
- Clearly define success criteria and acceptance criteria for each deliverable
- Quantify expected outcomes where possible (performance improvements, cost savings, efficiency gains)
- Emphasize business value and ROI rather than technical features
- Use client-focused language ("You will receive...", "This will enable you to...")
- Include timelines for when each deliverable will be completed
- Address how deliverables align with the client's strategic objectives
"""
        elif "pricing" in section.key.lower() or "investment" in section.key.lower():
            section_guidance = """
SPECIFIC GUIDANCE FOR PRICING/INVESTMENT SECTION:
- Present pricing as an investment in business outcomes, not just costs
- Break down pricing into clear, logical components
- Justify pricing with value proposition and ROI analysis
- Compare investment to potential business benefits and cost savings
- Provide flexible pricing options if appropriate (phases, modules, tiers)
- Address total cost of ownership, not just initial investment
- Include what is and isn't included in the pricing
- Use confident, value-focused language about the investment
"""
        elif "technical" in section.key.lower():
            section_guidance = """
SPECIFIC GUIDANCE FOR TECHNICAL SECTIONS:
- Only include technical details AFTER business value has been established
- Focus on how technical choices support business outcomes
- Avoid overwhelming technical jargon - keep it accessible to business stakeholders
- Emphasize proven technologies and industry standards
- Connect technical features to business benefits
"""

        prompt = f"""
You are an expert proposal writer. Generate comprehensive content for a single section of a professional proposal responding to an RFP. The response must be in {request.language}.

RFP CONTENT (for context):
{rfp_text[:4000]}

COMPANY DETAILS:
- Company: {request.company_name}
- Sector: {request.sector} 

{insights_context}

Generate detailed content for this specific section, in {request.language}:
{section_info}

{section_guidance}

GENERAL REQUIREMENTS:
1. The content for this section should be 500-800 words.
2. Address the specific RFP requirements and challenges related to this section.
3. Include specific examples, data, and case studies where relevant.
4. Ensure content directly responds to the RFP's needs for this topic.
5. Make the content engaging, persuasive, and highly detailed.
6. Use industry-specific terminology appropriately for the {request.language} language.
7. Structure the content with paragraphs, bullet points, and subheadings for readability.
8. IMPORTANT: Incorporate relevant insights from the special and additional documents where applicable to this section.
9. Reference standards, best practices, and methodologies from the supporting documents when relevant.
10. CRITICAL: For deliverables and pricing sections, focus on BUSINESS VALUE and CLIENT OUTCOMES first.

Respond with ONLY the detailed content for the section as a single string, in {request.language}. Do not wrap it in JSON or markdown.
"""

        try:
            response = ai_config.model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.7,
                    top_k=40,
                    top_p=0.95,
                    max_output_tokens=4000,
                )
            )
            return response.text.strip() if response.text else None
        except Exception as e:
            print(f"AI content generation for section '{section.key}' failed: {e}")
            return f"Content generation for '{section.title}' failed. Error: {e}"
    
    def _generate_mock_content(self, sections: List[Section], request: ProposalRequest) -> dict:
        content = {}
        
        for section in sections:
            content[section.key] = f"""This section addresses {section.title.lower()} for {request.company_name}. 
            
Our comprehensive approach ensures that all requirements are met through industry-leading practices and proven methodologies. We understand the critical importance of {section.title.lower()} in delivering successful project outcomes.

Key aspects include:
- Thorough analysis and planning
- Best practice implementation
- Quality assurance measures
- Continuous monitoring and improvement
- Detailed strategies for execution and delivery
- In-depth risk analysis and mitigation plans
- Advanced technical solutions and architectural designs
- Comprehensive timelines and resource allocation
- Clear communication and reporting protocols

This section is designed to be exceptionally detailed, providing over 1000 words of in-depth analysis and planning to fully address the requirements of the RFP. Our commitment to excellence is reflected in the depth of information provided, ensuring you have complete confidence in our ability to deliver.

{'Content requirements: ' + ', '.join(section.content_requirements) if section.content_requirements else ''}

We are committed to delivering exceptional results that exceed expectations and provide lasting value to your organization."""
        
        return content

    async def analyze_special_document(self, special_text: str, special_structure: ExtractedStructure, proposal_type: str, sector: str) -> str:
        """Analyze a special document to extract insights that will enhance the main proposal"""
        try:
            if ai_config.model:
                prompt = f"""
Analyze this special supporting document to extract key insights that will enhance a {proposal_type} proposal for the {sector} sector.

SPECIAL DOCUMENT CONTENT:
{special_text[:4000]}

EXTRACTED SECTIONS:
{json.dumps([s for s in special_structure.sections[:10]], indent=2)}

EXTRACTED REQUIREMENTS:
{json.dumps(special_structure.requirements[:10], indent=2)}

SCOPE:
{special_structure.scope}

Please provide key insights, best practices, methodologies, standards, or specific requirements from this document that should be incorporated into the main proposal. Focus on:

1. Technical standards and methodologies
2. Industry best practices
3. Compliance requirements
4. Quality standards
5. Implementation approaches
6. Risk mitigation strategies
7. Success metrics and KPIs

Respond with a concise but comprehensive analysis (maximum 1500 words) that can be used to enhance the main proposal content.
"""

                try:
                    response = ai_config.model.generate_content(prompt)
                    insights = response.text.strip()
                    print(f"Generated special document insights: {len(insights)} characters")
                    return insights
                except Exception as e:
                    print(f"Error generating special document insights with AI: {e}")
                    return self._generate_fallback_special_insights(special_text, proposal_type, sector)
            else:
                return self._generate_fallback_special_insights(special_text, proposal_type, sector)
                
        except Exception as e:
            print(f"Error analyzing special document: {e}")
            return self._generate_fallback_special_insights(special_text, proposal_type, sector)

    async def analyze_additional_documents(self, combined_text: str, proposal_type: str, sector: str) -> str:
        """Analyze additional supporting documents to extract insights for proposal enhancement"""
        try:
            if ai_config.model:
                prompt = f"""
Analyze these additional supporting documents to extract valuable insights for a {proposal_type} proposal in the {sector} sector.

ADDITIONAL DOCUMENTS CONTENT:
{combined_text[:5000]}

Please extract and synthesize key information that will enhance the main proposal, including:

1. Supporting evidence and case studies
2. Technical specifications and requirements
3. Industry standards and regulations
4. Best practices and methodologies  
5. Historical data and benchmarks
6. Stakeholder requirements and preferences
7. Implementation examples and lessons learned
8. Success metrics and evaluation criteria

Provide a well-organized analysis (maximum 2000 words) that identifies the most valuable insights to incorporate into the main proposal. Focus on actionable information that will strengthen the proposal's credibility and effectiveness.
"""

                try:
                    response = ai_config.model.generate_content(prompt)
                    insights = response.text.strip()
                    print(f"Generated additional documents insights: {len(insights)} characters")
                    return insights
                except Exception as e:
                    print(f"Error generating additional documents insights with AI: {e}")
                    return self._generate_fallback_additional_insights(combined_text, proposal_type, sector)
            else:
                return self._generate_fallback_additional_insights(combined_text, proposal_type, sector)
                
        except Exception as e:
            print(f"Error analyzing additional documents: {e}")
            return self._generate_fallback_additional_insights(combined_text, proposal_type, sector)

    def _generate_fallback_special_insights(self, special_text: str, proposal_type: str, sector: str) -> str:
        """Generate fallback insights when AI is not available"""
        insights = []
        
        # Extract key terms and phrases
        text_lower = special_text.lower()
        
        if 'standard' in text_lower or 'specification' in text_lower:
            insights.append("• Document contains technical standards and specifications that should be referenced in the technical approach section.")
        
        if 'requirement' in text_lower or 'must' in text_lower or 'shall' in text_lower:
            insights.append("• Document outlines specific requirements that must be addressed in the proposal solution.")
        
        if 'quality' in text_lower or 'compliance' in text_lower:
            insights.append("• Quality assurance and compliance considerations are highlighted that should be incorporated into the quality management section.")
        
        if 'risk' in text_lower or 'mitigation' in text_lower:
            insights.append("• Risk management strategies and mitigation approaches are identified that can strengthen the risk management section.")
        
        if 'timeline' in text_lower or 'schedule' in text_lower:
            insights.append("• Timeline and scheduling information is provided that should inform the project timeline and implementation plan.")
        
        if sector.lower() == 'health':
            if 'patient' in text_lower or 'clinical' in text_lower:
                insights.append("• Healthcare-specific considerations and patient-centered approaches are outlined.")
        elif sector.lower() == 'technology':
            if 'security' in text_lower or 'data' in text_lower:
                insights.append("• Technology security and data management requirements are specified.")
        
        return '\n'.join(insights) if insights else "Special document provides supplementary context that enhances understanding of project requirements and industry standards."

    def _generate_fallback_additional_insights(self, combined_text: str, proposal_type: str, sector: str) -> str:
        """Generate fallback insights for additional documents when AI is not available"""
        insights = []
        
        text_lower = combined_text.lower()
        
        # Count key terms to understand document focus
        technical_terms = ['technical', 'specification', 'system', 'solution', 'implementation']
        financial_terms = ['budget', 'cost', 'price', 'financial', 'funding']
        management_terms = ['project', 'management', 'team', 'coordination', 'leadership']
        
        technical_count = sum(text_lower.count(term) for term in technical_terms)
        financial_count = sum(text_lower.count(term) for term in financial_terms)
        management_count = sum(text_lower.count(term) for term in management_terms)
        
        if technical_count > 5:
            insights.append("• Documents contain significant technical content that should inform the technical approach and solution architecture.")
        
        if financial_count > 3:
            insights.append("• Financial information and cost considerations are provided that should be reflected in the budget and pricing sections.")
        
        if management_count > 3:
            insights.append("• Project management methodologies and team structure information are available to enhance the management approach.")
        
        if 'success' in text_lower or 'metric' in text_lower or 'kpi' in text_lower:
            insights.append("• Success metrics and key performance indicators are defined that should be incorporated into the evaluation criteria.")
        
        if 'experience' in text_lower or 'case study' in text_lower:
            insights.append("• Relevant experience and case studies are documented that can support the team qualifications and past performance sections.")
        
        return '\n'.join(insights) if insights else "Additional documents provide valuable supporting information and context that enhances the overall proposal quality and comprehensiveness."

# Enhanced Document Generator with Arabic Support Check
class EnhancedDocumentGenerator:
    def __init__(self):
        pass

    def _get_image_for_doc(self, path: Optional[str]) -> Optional[Union[str, BytesIO]]:
        print(f"_get_image_for_doc called with path: {path}")
        if not path:
            print("Path is None, returning None")
            return None

        if path.startswith(('http://', 'https://')):
            print("Path is URL, downloading...")
            try:
                response = requests.get(path, timeout=10)
                response.raise_for_status()
                return BytesIO(response.content)
            except requests.exceptions.RequestException as e:
                print(f"Error downloading image from URL {path}: {e}")
                return None
        
        print(f"Checking if file exists: {os.path.exists(path)}")
        if os.path.exists(path):
            print(f"File exists, returning path: {path}")
            return path
        
        print(f"Warning: Logo file not found for path: {path}")
        return None

    def _process_arabic_text(self, text):
        if not text or not ARABIC_SUPPORT:
            return text
        try:
            reshaped_text = arabic_reshaper.reshape(text)
            bidi_text = get_display(reshaped_text)
            return bidi_text
        except Exception as e:
            print(f"Arabic text processing error: {e}")
            return text

    def add_page_numbers_and_logos_to_word(self, doc, logo_top_left_path: Optional[str] = None, logo_bottom_right_path: Optional[str] = None, lang: str = 'en'):
        """Adds page numbers and logos to each section footer in a Word document."""
        
        print(f"Adding logos to document...")
        print(f"Top-left logo path: {logo_top_left_path}")
        print(f"Bottom-right logo path: {logo_bottom_right_path}")
        if logo_top_left_path:
            print(f"Top-left logo exists: {os.path.exists(logo_top_left_path)}")
        if logo_bottom_right_path:
            print(f"Bottom-right logo exists: {os.path.exists(logo_bottom_right_path)}")
        
        logo_top_left = self._get_image_for_doc(logo_top_left_path)
        logo_bottom_right = self._get_image_for_doc(logo_bottom_right_path)
        
        print(f"_get_image_for_doc returned - top_left: {logo_top_left}, bottom_right: {logo_bottom_right}")
        
        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            footer = section.footer
            footer.is_linked_to_previous = False

            if header.paragraphs:
                header.paragraphs[0].clear()
            if footer.paragraphs:
                footer.paragraphs[0].clear()

            if logo_top_left:
                header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = header_paragraph.add_run()
                try:
                    # Process image with PIL to ensure compatibility
                    from PIL import Image as PILImage
                    import io
                    
                    if isinstance(logo_top_left, str):
                        # Open and convert image to ensure compatibility
                        with PILImage.open(logo_top_left) as img:
                            # Convert to RGB if necessary (removes alpha channel issues)
                            if img.mode in ('RGBA', 'LA'):
                                img = img.convert('RGB')
                            
                            # Save to BytesIO as PNG (universally supported)
                            img_io = io.BytesIO()
                            img.save(img_io, format='PNG', optimize=True)
                            img_io.seek(0)
                            
                            run.add_picture(img_io, width=Inches(0.75))
                    else:
                        run.add_picture(logo_top_left, width=Inches(0.75))
                    if isinstance(logo_top_left, BytesIO):
                        logo_top_left.seek(0)
                    print("✅ Successfully added top-left logo to header")
                except Exception as e:
                    print(f"❌ Could not add top-left logo to header: {e}")
                    traceback.print_exc()

            footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            
            table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
            table.autofit = False
            table.columns[0].width = Inches(2.5)
            table.columns[1].width = Inches(1.5)
            table.columns[2].width = Inches(2.5)
            
            page_num_paragraph = table.cell(0, 1).paragraphs[0]
            page_num_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_page_number_field(page_num_paragraph, lang)
            
            if logo_bottom_right:
                logo_paragraph = table.cell(0, 2).paragraphs[0]
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = logo_paragraph.add_run()
                try:
                    # Process image with PIL to ensure compatibility
                    from PIL import Image as PILImage
                    import io
                    
                    if isinstance(logo_bottom_right, str):
                        # Open and convert image to ensure compatibility
                        with PILImage.open(logo_bottom_right) as img:
                            # Convert to RGB if necessary (removes alpha channel issues)
                            if img.mode in ('RGBA', 'LA'):
                                img = img.convert('RGB')
                            
                            # Save to BytesIO as PNG (universally supported)
                            img_io = io.BytesIO()
                            img.save(img_io, format='PNG', optimize=True)
                            img_io.seek(0)
                            
                            run.add_picture(img_io, width=Inches(0.75))
                    else:
                        run.add_picture(logo_bottom_right, width=Inches(0.75))
                    if isinstance(logo_bottom_right, BytesIO):
                        logo_bottom_right.seek(0)
                    print("✅ Successfully added bottom-right logo to footer")
                except Exception as e:
                    print(f"❌ Could not add bottom-right logo to footer: {e}")
                    traceback.print_exc()

    def _add_bookmark(self, paragraph, bookmark_name):
        """Add a bookmark to a paragraph"""
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        
        # Create bookmark start
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), '0')
        bookmark_start.set(qn('w:name'), bookmark_name)
        run._r.insert(0, bookmark_start)
        
        # Create bookmark end
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), '0')
        run._r.append(bookmark_end)

    def _add_page_ref_field(self, paragraph, bookmark_name):
        """Add a page reference field that shows the page number of a bookmark"""
        run = paragraph.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar_begin)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f' PAGEREF {bookmark_name} \\h '
        run._r.append(instrText)
        
        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar_sep)
        
        t = OxmlElement('w:t')
        t.text = '1'  # Default text before field is updated
        run._r.append(t)

        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar_end)

    def _add_field_update_instruction(self, doc):
        """Add an instruction to update all fields when document is opened"""
        # This is a simple approach - add a hidden instruction paragraph
        # Word will show the instruction when fields need updating
        settings_paragraph = doc.add_paragraph()
        settings_run = settings_paragraph.add_run()
        
        # Add an UpdateFields instruction (this will prompt Word to update fields)
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        settings_run._r.append(fldChar_begin)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' DOCPROPERTY "UpdateFields" \\* MERGEFORMAT '
        settings_run._r.append(instrText)
        
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        settings_run._r.append(fldChar_end)
        
        # Hide this paragraph
        settings_paragraph.paragraph_format.hidden = True

    def _add_page_number_field(self, paragraph, lang: str = 'en'):
        """Adds 'Page X of Y' field to a paragraph."""
        
        def create_field(p, field_code):
            run = p.add_run()
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar_begin)
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = field_code
            run._r.append(instrText)
            
            fldChar_sep = OxmlElement('w:fldChar')
            fldChar_sep.set(qn('w:fldCharType'), 'separate')
            run._r.append(fldChar_sep)
            
            t = OxmlElement('w:t')
            t.text = '?'
            run._r.append(t)

            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar_end)

        page_text = TRANSLATIONS[lang]['page']
        of_text = TRANSLATIONS[lang]['of']

        paragraph.add_run(f'{page_text} ')
        create_field(paragraph, 'PAGE')
        paragraph.add_run(f' {of_text} ')
        create_field(paragraph, 'NUMPAGES')
    
    def generate_word_document(self, content: dict, structure: List[Section], company_name: str, job_id: str, request: ProposalRequest) -> str:
        """Generate Word document with dynamic structure and TOC"""
        try:
            lang = request.language if request.language in TRANSLATIONS else 'en'
            doc = Document()

            if lang == 'ar':
                doc_element = doc.element.find(qn('w:body'))
                if doc_element is not None:
                    sectPr = doc_element.find(qn('w:sectPr'))
                    if sectPr is not None:
                        bidi = sectPr.find(qn('w:bidi'))
                        if bidi is None:
                            bidi = OxmlElement('w:bidi')
                            sectPr.append(bidi)
                        bidi.set(qn('w:val'), 'on')
            
            self._create_custom_styles(doc)
            self._create_title_page(doc, company_name, lang)
            self._create_dynamic_toc(doc, structure, request, lang)
            self._add_dynamic_content(doc, structure, content, request, lang)
            self.add_page_numbers_and_logos_to_word(
                doc, 
                request.logo_top_left_path, 
                request.logo_bottom_right_path, 
                lang
            )
            
            # Add a field update instruction
            self._add_field_update_instruction(doc)
            
            filename = f"proposal_{company_name.replace(' ', '_')}_{job_id}.docx"
            filepath = os.path.join(OUTPUT_DIR, filename)
            doc.save(filepath)
            
            return filename
            
        except Exception as e:
            print(f"Error generating Word document: {e}")
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"Error generating Word document: {str(e)}")
    
    def _create_custom_styles(self, doc):
        """Create custom styles for the document"""
        styles = doc.styles
        
        for i in range(1, 5):
            style_name = f'CustomHeading{i}'
            if style_name not in styles:
                style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                style.base_style = styles[f'Heading {i}']
                style.font.name = 'Arial'
                style.font.size = Pt(20 - (i * 2))
                style.font.bold = True
                style.font.color.rgb = RGBColor(0, 51, 102)
    
    def _create_title_page(self, doc, company_name: str, lang: str):
        """Create professional title page"""
        title_text = TRANSLATIONS[lang]['technical_proposal']
        title = doc.add_heading(title_text, 0)
        
        prepared_for_text = f"{TRANSLATIONS[lang]['prepared_for']} {company_name}"
        subtitle = doc.add_paragraph(prepared_for_text)

        if lang == 'ar':
            title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            subtitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in title.runs:
                run.font.rtl = True
            for run in subtitle.runs:
                run.font.rtl = True
        else:
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].font.bold = True
        
        doc.add_paragraph()
        
        date_para = doc.add_paragraph(datetime.now().strftime('%B %Y'))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if lang == 'en' else WD_ALIGN_PARAGRAPH.RIGHT
        date_para.runs[0].font.size = Pt(14)
        
        doc.add_page_break()
    
    def _create_dynamic_toc(self, doc, structure: List[Section], request: ProposalRequest, lang: str):
        """Create table of contents based on dynamic structure"""
        toc_heading_text = TRANSLATIONS[lang]['table_of_contents']
        toc_heading = doc.add_heading(toc_heading_text, level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT if lang == 'en' else WD_ALIGN_PARAGRAPH.RIGHT
        if lang == 'ar':
            for run in toc_heading.runs:
                run.font.rtl = True

        # Add instruction for updating page numbers
        instruction_p = doc.add_paragraph()
        instruction_text = "Note: To update page numbers, select all text (Ctrl+A) and press F9, or right-click and select 'Update Field'."
        instruction_run = instruction_p.add_run(instruction_text)
        instruction_run.font.size = Pt(9)
        instruction_run.font.italic = True
        instruction_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        instruction_p.paragraph_format.space_after = Pt(12)

        self._add_toc_entries_recursive(doc, structure, request, lang)
        doc.add_page_break()
    
    def _add_toc_entries_recursive(self, doc, sections: List[Section], request: ProposalRequest, lang: str, level: int = 0):
        """Recursively add TOC entries with proper page numbers"""
        for section in sections:
            if not self._should_include_section(section, request):
                continue
            
            p = doc.add_paragraph()
            if lang == 'en':
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                indent = '    ' * level
                
                # Add section title (clean from * and # characters)
                clean_title = section.title.replace('*', '').replace('#', '')
                run1 = p.add_run(f"{indent}{section.number}. {clean_title}")
                
                # Add dots to fill space
                run2 = p.add_run()
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.0))
                run2.add_tab()
                
                # Add page number field
                self._add_page_ref_field(p, f"section_{section.key}")
                
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(f"{section.title} .{section.number}")
                run.font.rtl = True

            p.paragraph_format.space_after = Pt(6)
            
            if section.subsections:
                self._add_toc_entries_recursive(doc, section.subsections, request, lang, level + 1)
    
    def _add_dynamic_content(self, doc, sections: List[Section], content: dict, request: ProposalRequest, lang: str):
        """Add content based on dynamic structure"""
        for section in sections:
            if not self._should_include_section(section, request):
                continue
            
            # Clean title by removing * and # characters
            clean_title = section.title.replace('*', '').replace('#', '')
            heading_text = f'{section.number}. {clean_title}'
            heading = doc.add_heading(heading_text, level=min(section.level, 4))
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT if lang == 'en' else WD_ALIGN_PARAGRAPH.RIGHT
            if lang == 'ar':
                for run in heading.runs:
                    run.font.rtl = True

            # Add bookmark for TOC page reference
            self._add_bookmark(heading, f"section_{section.key}")

            if section.key in content and content[section.key]:
                # Clean content by removing * and # characters
                cleaned_content = content[section.key].replace('*', '').replace('#', '')
                lines = cleaned_content.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    if line.startswith(('- ', '• ')):
                        p = doc.add_paragraph(line[2:], style='List Bullet')
                    else:
                        p = doc.add_paragraph(line)
                    
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if lang == 'en' else WD_ALIGN_PARAGRAPH.RIGHT
                    if lang == 'ar':
                        for run in p.runs:
                            run.font.rtl = True
                    p.paragraph_format.space_after = Pt(12)
            
            doc.add_paragraph()
            
            if section.subsections:
                self._add_dynamic_content(doc, section.subsections, content, request, lang)
            
            if section.level == 1 and section != sections[-1]:
                doc.add_page_break()
    
    def _should_include_section(self, section: Section, request: ProposalRequest) -> bool:
        """Check if section should be included based on user selection"""
        if not request.selected_sections:
            return True
        return section.key in request.selected_sections
    
    def generate_pdf_document(self, content: dict, structure: List[Section], company_name: str, job_id: str, request: ProposalRequest) -> str:
        """Generate PDF document with dynamic structure"""
        try:
            lang = request.language if request.language in TRANSLATIONS else 'en'
            
            # Check for Arabic support and font availability
            if lang == 'ar':
                if not ARABIC_SUPPORT:
                    raise Exception(
                        "Cannot generate Arabic PDF because Arabic text processing libraries are not installed. "
                        "Please install python-bidi and arabic-reshaper packages."
                    )
                if not AMIRI_FONT_PATH:
                    raise Exception(
                        "Cannot generate Arabic PDF because the Amiri font is missing or failed to load. "
                        "Please ensure 'Amiri-Regular.ttf' is in the 'backend/fonts/' directory."
                    )

            filename = f"proposal_{company_name.replace(' ', '_')}_{job_id}.pdf"
            filepath = os.path.join(OUTPUT_DIR, filename)
            
            doc = SimpleDocTemplate(filepath, pagesize=A4)
            styles = getSampleStyleSheet()

            if lang == 'ar' and AMIRI_FONT_PATH:
                styles.add(ParagraphStyle(name='ArabicTitle', fontName='AmiriFamily', fontSize=24, alignment=2, spaceAfter=30, textColor=colors.HexColor('#003366')))
                styles.add(ParagraphStyle(name='ArabicSubtitle', fontName='AmiriFamily', fontSize=16, alignment=2, spaceAfter=20, textColor=colors.HexColor('#336699')))
                styles.add(ParagraphStyle(name='ArabicHeading1', fontName='AmiriFamily', fontSize=18, alignment=2, spaceAfter=12))
                styles.add(ParagraphStyle(name='ArabicHeading2', fontName='AmiriFamily', fontSize=14, alignment=2, spaceAfter=10))
                styles.add(ParagraphStyle(name='ArabicHeading3', fontName='AmiriFamily', fontSize=12, alignment=2, spaceAfter=8))
                styles.add(ParagraphStyle(name='ArabicBody', fontName='AmiriFamily', fontSize=10, alignment=2, leading=14))
                styles.add(ParagraphStyle(name='ArabicTOC', fontName='AmiriFamily', fontSize=10, alignment=2, rightIndent=20))

            story = []

            title_style = styles['Title'] if lang == 'en' else styles['ArabicTitle']
            subtitle_style = styles['Heading2'] if lang == 'en' else styles['ArabicSubtitle']
            normal_style = styles['Normal'] if lang == 'en' else styles['ArabicBody']

            title_text = TRANSLATIONS[lang]['technical_proposal']
            story.append(Paragraph(self._process_arabic_text(title_text) if lang == 'ar' else title_text, title_style))
            story.append(Spacer(1, 20))

            prepared_for_text = f"{TRANSLATIONS[lang]['prepared_for']} {company_name}"
            story.append(Paragraph(self._process_arabic_text(prepared_for_text) if lang == 'ar' else prepared_for_text, subtitle_style))
            story.append(Spacer(1, 10))
            story.append(Paragraph(datetime.now().strftime('%B %Y'), normal_style))
            story.append(PageBreak())
            
            toc_heading_text = TRANSLATIONS[lang]['table_of_contents']
            story.append(Paragraph(self._process_arabic_text(toc_heading_text) if lang == 'ar' else toc_heading_text, styles['Heading1'] if lang == 'en' else styles['ArabicHeading1']))
            story.append(Spacer(1, 20))
            
            self._add_pdf_toc_recursive(story, structure, styles, request, lang)
            story.append(PageBreak())
            
            self._add_pdf_content_recursive(story, structure, content, styles, request, lang)

            buffer = BytesIO()
            doc_for_count = SimpleDocTemplate(buffer, pagesize=A4)
            doc_for_count.build(story)
            total_pages = doc_for_count.page
            buffer.close()

            # Pass 2: Build the final PDF with page numbers
            handler = lambda canvas, doc: self._add_page_numbers_and_logos_to_pdf(canvas, doc, request.logo_top_left_path, request.logo_bottom_right_path, total_pages, lang)
            doc.build(story, onFirstPage=handler, onLaterPages=handler)
            
            return filename
            
        except Exception as e:
            print(f"Error generating PDF document: {e}")
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"Error generating PDF document: {str(e)}")

    def _add_page_numbers_and_logos_to_pdf(self, canvas, doc, logo_top_left_path, logo_bottom_right_path, total_pages, lang: str = 'en'):
        """Add page numbers and logos to each page of a PDF document."""
        logo_top_left = self._get_image_for_doc(logo_top_left_path)
        logo_bottom_right = self._get_image_for_doc(logo_bottom_right_path)
        canvas.saveState()
        
        page_text = TRANSLATIONS[lang]['page']
        of_text = TRANSLATIONS[lang]['of']
        page_number_text = f"{page_text} {doc.page} {of_text} {total_pages}"
        
        if lang == 'ar' and AMIRI_FONT_PATH:
            canvas.setFont('AmiriFamily', 9)
            page_number_text = self._process_arabic_text(page_number_text)
        else:
            canvas.setFont('Times-Roman', 9)

        canvas.drawCentredString(letter[0] / 2, 0.75 * inch, page_number_text)
        
        if logo_top_left:
            try:
                canvas.drawImage(logo_top_left, 0.5 * inch, A4[1] - 1 * inch, width=0.5 * inch, height=0.5 * inch, mask='auto')
                if isinstance(logo_top_left, BytesIO):
                    logo_top_left.seek(0)
            except Exception as e:
                print(f"Could not add top-left logo to PDF: {e}")

        if logo_bottom_right:
            try:
                canvas.drawImage(logo_bottom_right, A4[0] - 1 * inch, 0.5 * inch, width=0.5 * inch, height=0.5 * inch, mask='auto')
                if isinstance(logo_bottom_right, BytesIO):
                    logo_bottom_right.seek(0)
            except Exception as e:
                print(f"Could not add bottom-right logo to PDF: {e}")

        canvas.restoreState()
    
    def _add_pdf_toc_recursive(self, story: list, sections: List[Section], styles, request: ProposalRequest, lang: str, level: int = 0):
        """Add PDF table of contents entries recursively"""
        for section in sections:
            if not self._should_include_section(section, request):
                continue
            
            if lang == 'en':
                toc_style = styles['Normal']
                indent = '    ' * level
                toc_text = f"{indent}{section.number}. {section.title}"
            else:
                toc_style = styles['ArabicTOC']
                toc_text = self._process_arabic_text(f"{section.title} .{section.number}")
            
            story.append(Paragraph(toc_text, toc_style))
            story.append(Spacer(1, 6))
            
            if section.subsections:
                self._add_pdf_toc_recursive(story, section.subsections, styles, request, lang, level + 1)
    
    def _add_pdf_content_recursive(self, story: list, sections: List[Section], content: dict, styles, request: ProposalRequest, lang: str):
        """Add PDF content recursively"""
        for section in sections:
            if not self._should_include_section(section, request):
                continue
            
            heading_level = min(section.level, 3)
            if lang == 'en':
                heading_style = styles[f'Heading{heading_level}']
                heading_text = f'{section.number}. {section.title}'
            else:
                heading_style = styles[f'ArabicHeading{heading_level}']
                heading_text = self._process_arabic_text(f'{section.number}. {section.title}')

            story.append(Paragraph(heading_text, heading_style))
            story.append(Spacer(1, 12))
            
            if section.key in content and content[section.key]:
                # Clean content by removing * and # characters
                cleaned_content = content[section.key].replace('*', '').replace('#', '')
                body_style = styles['Normal'] if lang == 'en' else styles['ArabicBody']
                paragraphs = cleaned_content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        if para_text.strip().startswith(('* ', '- ')):
                            p_text = f"• {para_text.strip()[2:]}"
                        else:
                            p_text = para_text.strip()
                        
                        processed_text = self._process_arabic_text(p_text) if lang == 'ar' else p_text
                        story.append(Paragraph(processed_text, body_style))
                        story.append(Spacer(1, 10))
            
            if section.subsections:
                self._add_pdf_content_recursive(story, section.subsections, content, styles, request, lang)
            
            if section.level == 1:
                story.append(Spacer(1, 20))
                story.append(PageBreak())

# Global instances
processor = EnhancedDocumentProcessor()
ai_generator = EnhancedAIContentGenerator()
doc_generator = EnhancedDocumentGenerator()
diagram_generator = MermaidDiagramGenerator()
visualization_generator = VisualizationDocumentGenerator()

# Storage for job status and generated structures
job_status = {}
generated_structures = {}
active_jobs = set()  # Track active job IDs to prevent duplicates

# Enhanced API endpoints
@app.post("/test-upload")
async def test_upload(
    files: List[UploadFile] = File(...),
    proposal_type: str = Form(...),
    sector: str = Form(...),
    company_name: str = Form(...),
    special_document: Optional[UploadFile] = File(None),
    additional_documents: List[UploadFile] = File(default=[]),
    language: Optional[str] = Form("en")
):
    """Simple test endpoint to debug form submission issues"""
    return {
        "status": "success",
        "received": {
            "proposal_type": proposal_type,
            "sector": sector,
            "company_name": company_name,
            "language": language,
            "files_count": len(files),
            "special_document": special_document.filename if special_document else None,
            "additional_documents_count": len(additional_documents)
        }
    }

@app.get("/")
async def root():
    return {
        "message": "Enhanced RFP Proposal Generator API v2.1 - Dynamic Structure Generation",
        "features": [
            "Dynamic table of contents generation from RFP analysis",
            "AI-powered structure extraction and proposal generation", 
            "Intelligent section hierarchy based on RFP content",
            "Customizable section selection and filtering",
            "Professional document formatting with proper TOC",
            "Support for multiple document formats (PDF, DOCX)",
            "Arabic language support with proper font handling"
        ],
        "arabic_support": ARABIC_SUPPORT and AMIRI_FONT_PATH is not None
    }

@app.post("/analyze-rfp")
async def analyze_rfp_structure(
    files: List[UploadFile] = File(...),
    company_name: str = Form(...),
    sector: str = Form("general")
):
    """Analyze uploaded RFP and return suggested proposal structure"""
    try:
        combined_text = ""
        all_structures = []
        
        temp_files = []
        for file in files:
            temp_path = os.path.join(UPLOAD_DIR, f"temp_{uuid.uuid4()}_{file.filename}")
            temp_files.append(temp_path)
            
            with open(temp_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            text, structure = processor.process_file(temp_path)
            combined_text += text + "\n\n"
            all_structures.append(structure)
        
        for temp_path in temp_files:
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
        combined_structure = ExtractedStructure(
            sections=[],
            requirements=[],
            scope=""
        )
        
        for struct in all_structures:
            combined_structure.sections.extend(struct.sections)
            combined_structure.requirements.extend(struct.requirements)
            combined_structure.scope += " " + struct.scope
        
        proposal_structure = await ai_generator.analyze_rfp_and_generate_structure(
            combined_text, combined_structure, ProposalRequest(
                proposal_type="technical",
                sector="general",
                company_name=company_name,
                selected_sections=None,
                output_format="all",
                logo_top_left_path=None,
                logo_bottom_right_path=None,
                language="en"
            )
        )
        
        return {
            "success": True,
            "extracted_info": {
                "sections_found": len(combined_structure.sections),
                "requirements_found": len(combined_structure.requirements),
                "scope_summary": combined_structure.scope[:500] + "..." if len(combined_structure.scope) > 500 else combined_structure.scope
            },
            "suggested_structure": [section.to_dict() for section in proposal_structure],
            "available_sections": [
                {
                    "key": section.key,
                    "title": section.title,
                    "level": section.level,
                    "content_requirements": section.content_requirements
                }
                for section in ai_generator.flatten_sections(proposal_structure)
            ]
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error analyzing RFP: {str(e)}")

@app.post("/upload-and-generate", response_model=ProposalResponse)
async def upload_and_generate_enhanced(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    proposal_type: str = Form(...),
    sector: str = Form(...),
    company_name: str = Form(...),
    selected_sections: Optional[str] = Form(None),
    use_dynamic_structure: Optional[bool] = Form(True),
    output_format: Optional[str] = Form("all"),
    logo_top_left: Optional[UploadFile] = File(None),
    logo_bottom_right: Optional[UploadFile] = File(None),
    special_document: Optional[UploadFile] = File(None),
    additional_documents: List[UploadFile] = File(default=[]),
    language: Optional[str] = Form("en")
):
    """Enhanced upload and generate with dynamic structure"""
    try:
        print("--- New Upload Request ---")
        print(f"Received parameters:")
        print(f"  proposal_type: {proposal_type}")
        print(f"  sector: {sector}")
        print(f"  company_name: {company_name}")
        print(f"  language: {language}")
        print(f"  files count: {len(files) if files else 0}")
        print(f"  special_document: {special_document.filename if special_document else None}")
        print(f"  additional_documents: {len(additional_documents)}")
        if logo_top_left:
            print(f"Received top-left logo: {logo_top_left.filename}")
        else:
            print("No top-left logo received.")
        
        if logo_bottom_right:
            print(f"Received bottom-right logo: {logo_bottom_right.filename}")
        else:
            print("No bottom-right logo received.")
        print("--------------------------")

        job_id = str(uuid.uuid4())
        
        selected_sections_list = None
        if selected_sections:
            selected_sections_list = [s.strip() for s in selected_sections.split(',')]
        
        # Check if there are too many active jobs (prevent system overload)
        if len(active_jobs) > 10:
            raise HTTPException(status_code=429, detail="Too many active jobs. Please wait and try again.")
        
        # Add to active jobs
        active_jobs.add(job_id)
        
        job_status[job_id] = {
            "status": "processing",
            "message": "Analyzing RFP and generating proposal structure...",
            "progress": 10,
            "files": []
        }
        
        logo_top_left_path = None
        if logo_top_left:
            relative_path = os.path.join(UPLOAD_DIR, f"{job_id}_{logo_top_left.filename}")
            with open(relative_path, "wb") as buffer:
                shutil.copyfileobj(logo_top_left.file, buffer)
            logo_top_left_path = os.path.abspath(relative_path)
            print(f"Saved top-left logo to: {logo_top_left_path}")
            print(f"File exists: {os.path.exists(logo_top_left_path)}")

        logo_bottom_right_path = None
        if logo_bottom_right:
            relative_path = os.path.join(UPLOAD_DIR, f"{job_id}_{logo_bottom_right.filename}")
            with open(relative_path, "wb") as buffer:
                shutil.copyfileobj(logo_bottom_right.file, buffer)
            logo_bottom_right_path = os.path.abspath(relative_path)
            print(f"Saved bottom-right logo to: {logo_bottom_right_path}")
            print(f"File exists: {os.path.exists(logo_bottom_right_path)}")
        
        # Process special document if provided
        special_document_content = None
        special_document_insights = None
        if special_document:
            print(f"Processing special document: {special_document.filename}")
            try:
                special_temp_path = os.path.join(UPLOAD_DIR, f"{job_id}_special_{special_document.filename}")
                with open(special_temp_path, "wb") as buffer:
                    shutil.copyfileobj(special_document.file, buffer)
                
                # Extract content from special document
                special_text, special_structure = processor.process_file(special_temp_path)
                special_document_content = special_text
                
                # Generate insights from special document using AI
                special_document_insights = await ai_generator.analyze_special_document(
                    special_text, special_structure, proposal_type, sector
                )
                
                # Clean up temporary file
                if os.path.exists(special_temp_path):
                    os.remove(special_temp_path)
                    
                print(f"Special document processed successfully, extracted {len(special_text)} characters")
            except Exception as e:
                print(f"Error processing special document: {str(e)}")
            finally:
                special_document.file.close()
        
        # Process additional documents if provided
        additional_documents_content = []
        additional_documents_insights = None
        if additional_documents:
            print(f"Processing {len(additional_documents)} additional documents")
            try:
                for i, doc in enumerate(additional_documents):
                    doc_temp_path = os.path.join(UPLOAD_DIR, f"{job_id}_additional_{i}_{doc.filename}")
                    with open(doc_temp_path, "wb") as buffer:
                        shutil.copyfileobj(doc.file, buffer)
                    
                    # Extract content from additional document
                    doc_text, doc_structure = processor.process_file(doc_temp_path)
                    additional_documents_content.append(doc_text)
                    
                    # Clean up temporary file
                    if os.path.exists(doc_temp_path):
                        os.remove(doc_temp_path)
                    
                    doc.file.close()
                
                # Generate combined insights from all additional documents
                if additional_documents_content:
                    combined_additional_text = "\n\n".join(additional_documents_content)
                    additional_documents_insights = await ai_generator.analyze_additional_documents(
                        combined_additional_text, proposal_type, sector
                    )
                    
                print(f"Additional documents processed successfully, total content: {sum(len(text) for text in additional_documents_content)} characters")
            except Exception as e:
                print(f"Error processing additional documents: {str(e)}")
        
        request = ProposalRequest(
            proposal_type=proposal_type,
            sector=sector,
            company_name=company_name,
            selected_sections=selected_sections_list,
            output_format=output_format,
            logo_top_left_path=logo_top_left_path,
            logo_bottom_right_path=logo_bottom_right_path,
            language=language,
            special_document_content=special_document_content,
            additional_documents_content=additional_documents_content,
            special_document_insights=special_document_insights,
            additional_documents_insights=additional_documents_insights
        )
        
        temp_file_paths = []
        for file in files:
            temp_path = os.path.join(UPLOAD_DIR, f"{job_id}_{file.filename}")
            try:
                with open(temp_path, "wb") as buffer:
                    shutil.copyfileobj(file.file, buffer)
                temp_file_paths.append(temp_path)
            finally:
                file.file.close()

        background_tasks.add_task(
            process_enhanced_proposal, 
            job_id, 
            temp_file_paths, 
            request, 
            use_dynamic_structure
        )
        
        return ProposalResponse(
            job_id=job_id,
            status="processing",
            message="Your proposal is being generated with dynamic structure analysis."
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error starting proposal generation: {str(e)}")

async def process_enhanced_proposal(
    job_id: str, 
    temp_file_paths: List[str], 
    request: ProposalRequest, 
    use_dynamic_structure: bool = True
):
    """Enhanced background task with dynamic structure generation and proper error handling"""
    print(f"BACKGROUND TASK STARTED for job {job_id}")
    print(f"Logo paths in request: top={request.logo_top_left_path}, bottom={request.logo_bottom_right_path}")
    try:
        job_status[job_id]["message"] = "Processing uploaded RFP files..."
        job_status[job_id]["progress"] = 20

        # Process uploaded files
        combined_text = ""
        all_structures = []
        
        for temp_path in temp_file_paths:
            text, structure = processor.process_file(temp_path)
            combined_text += text + "\n\n"
            all_structures.append(structure)
        
        for temp_path in temp_file_paths:
            if os.path.exists(temp_path):
                os.remove(temp_path)

        job_status[job_id]["message"] = "Analyzing RFP structure and generating proposal outline..."
        job_status[job_id]["progress"] = 40

        combined_structure = ExtractedStructure(
            sections=[],
            requirements=[],
            scope=""
        )
        
        for struct in all_structures:
            combined_structure.sections.extend(struct.sections)
            combined_structure.requirements.extend(struct.requirements)
            combined_structure.scope += " " + struct.scope

        if use_dynamic_structure:
            proposal_structure = await ai_generator.analyze_rfp_and_generate_structure(
                combined_text, combined_structure, request
            )
        else:
            proposal_structure = ai_generator._generate_fallback_structure(combined_structure)
            ai_generator.number_sections(proposal_structure)
        
        generated_structures[job_id] = proposal_structure

        job_status[job_id]["message"] = "Generating proposal content..."
        job_status[job_id]["progress"] = 60

        # Generate content and files based on proposal type
        generated_files = []
        
        try:
            if request.proposal_type == "technical":
                content = await ai_generator.generate_proposal_content(
                    combined_text, proposal_structure, request
                )
                
                # Generate Word document
                if request.output_format in ["all", "docx"]:
                    try:
                        word_file = doc_generator.generate_word_document(
                            content, proposal_structure, request.company_name, job_id, request
                        )
                        generated_files = [word_file]
                    except Exception as e:
                        print(f"Error generating Word document: {e}")
                        job_status[job_id]["message"] = f"Warning: Word document generation failed: {str(e)}"
                
                # Generate Visualization HTML document (ALWAYS generated alongside main document)
                try:
                    job_status[job_id]["message"] = "Generating interactive visualizations..."
                    job_status[job_id]["progress"] = 85
                    
                    visualization_file = visualization_generator.generate_visualization_html(
                        content, proposal_structure, request.company_name, job_id, request
                    )
                    generated_files.append(visualization_file)
                    print(f"✅ Generated visualization file: {visualization_file}")
                except Exception as e:
                    print(f"⚠️  Warning: Visualization generation failed: {e}")
                    # Continue even if visualization fails - don't break main generation
                
                # Generate PDF document
                # if request.output_format in ["all", "pdf"]:
                #     try:
                #         pdf_file = doc_generator.generate_pdf_document(
                #             content, proposal_structure, request.company_name, job_id, request
                #         )
                #         generated_files.append(pdf_file)
                #     except Exception as e:
                #         print(f"Error generating PDF document: {e}")
                #         error_msg = str(e)
                #         if "Arabic" in error_msg or "Amiri" in error_msg or "bidi" in error_msg:
                #             job_status[job_id]["message"] = f"Warning: PDF generation failed - {error_msg}"
                #         else:
                #             job_status[job_id]["message"] = f"Warning: PDF document generation failed: {error_msg}"
                
            elif request.proposal_type == "financial":
                financial_content = await generate_financial_content(combined_text, request)
                excel_file = generate_excel_financial_enhanced(
                    financial_content, proposal_structure, request.company_name, job_id, request
                )
                generated_files = [excel_file]
                
                # Generate Visualization HTML document for financial proposals too
                try:
                    job_status[job_id]["message"] = "Generating financial visualizations..."
                    job_status[job_id]["progress"] = 85
                    
                    visualization_file = visualization_generator.generate_visualization_html(
                        financial_content, proposal_structure, request.company_name, job_id, request
                    )
                    generated_files.append(visualization_file)
                    print(f"✅ Generated financial visualization file: {visualization_file}")
                except Exception as e:
                    print(f"⚠️  Warning: Financial visualization generation failed: {e}")
                    # Continue even if visualization fails
            
        except Exception as e:
            print(f"Error in document generation phase: {e}")
            traceback.print_exc()
            if not generated_files:  # Only fail if no files were generated
                raise e

        job_status[job_id]["message"] = "Finalizing documents..."
        job_status[job_id]["progress"] = 90

        # Check if any files were actually generated
        if not generated_files:
            raise Exception("No files were successfully generated. Please check your language and format selections.")

        # Final status update
        success_message = f"Proposal generated successfully with {len(ai_generator.flatten_sections(proposal_structure))} sections!"
        if len(generated_files) == 1:
            success_message += f" Generated: {generated_files[0]}"
        else:
            success_message += f" Generated {len(generated_files)} files."

        job_status[job_id] = {
            "status": "completed",
            "message": success_message,
            "progress": 100,
            "files": generated_files,
            "structure_summary": {
                "total_sections": len(ai_generator.flatten_sections(proposal_structure)),
                "main_sections": len(proposal_structure),
                "dynamic_generation": use_dynamic_structure
            }
        }
        
        # Remove from active jobs
        active_jobs.discard(job_id)
        
        asyncio.create_task(schedule_cleanup(job_id, 300))

    except Exception as e:
        print("Error in process_enhanced_proposal:", e)
        traceback.print_exc()
        
        # Remove from active jobs on error
        active_jobs.discard(job_id)
        
        job_status[job_id] = {
            "status": "error",
            "message": f"Error generating proposal: {str(e)}",
            "progress": 0,
            "files": []
        }
        asyncio.create_task(schedule_cleanup(job_id, 300))

async def schedule_cleanup(job_id: str, delay_seconds: int):
    """Schedules a cleanup task to remove job data after a delay."""
    await asyncio.sleep(delay_seconds)
    cleanup_job_data(job_id)

def cleanup_job_data(job_id: str):
    """Cleans up the job data from the global dictionaries."""
    if job_id in job_status:
        files_to_delete = job_status[job_id].get("files", [])
        for filename in files_to_delete:
            file_path = os.path.join(OUTPUT_DIR, filename)
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                    print(f"Cleaned up file: {file_path}")
                except OSError as e:
                    print(f"Error cleaning up file {file_path}: {e}")

        del job_status[job_id]

    if job_id in generated_structures:
        del generated_structures[job_id]
    
    # Remove from active jobs if still there
    active_jobs.discard(job_id)
    
    print(f"Cleaned up job data for {job_id}")

async def generate_financial_content(rfp_text: str, request: ProposalRequest) -> dict:
    """Generate financial proposal content using AI"""
    if not ai_config.model:
        return {
            "financial_summary": "Financial proposal summary not available - AI model not configured",
            "payment_schedule": []
        }
    
    prompt = f'''
Based on this RFP content, generate a comprehensive financial proposal.

RFP CONTENT:
{rfp_text[:3000]}

COMPANY: {request.company_name}
SECTOR: {request.sector}

Generate a JSON object with:
- "financial_summary": Professional summary of the financial offer (200-400 words)
- "payment_schedule": Array of payment milestones with "phase", "description", "amount", "percent"
- "total_investment": Total project cost
- "roi_projection": Expected return on investment details

Example structure:
{{
  "financial_summary": "We are pleased to present our comprehensive financial proposal...",
  "payment_schedule": [
    {{"phase": "Phase 1", "description": "Project Initiation and Analysis", "amount": 150000, "percent": "25%"}},
    {{"phase": "Phase 2", "description": "Development and Implementation", "amount": 300000, "percent": "50%"}},
    {{"phase": "Phase 3", "description": "Testing and Deployment", "amount": 150000, "percent": "25%"}}
  ],
  "total_investment": 600000,
  "roi_projection": "Expected ROI of 200% within 24 months"
}}

Respond ONLY with valid JSON.
'''
    
    try:
        response = ai_config.model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.7,
                top_k=40,
                top_p=0.95,
                max_output_tokens=4000,
            )
        )
        
        if response.text:
            cleaned_text = response.text.strip()
            if cleaned_text.startswith('```json'):
                cleaned_text = cleaned_text.replace('```json', '').replace('```', '').strip()
            
            return json.loads(cleaned_text)
    
    except Exception as e:
        print(f"Financial content generation failed: {e}")
    
    return {
        "financial_summary": f"Comprehensive financial proposal for {request.company_name} addressing all RFP requirements with competitive pricing and flexible payment terms.",
        "payment_schedule": [
            {"phase": "Phase 1", "description": "Project Initiation", "amount": 100000, "percent": "30%"},
            {"phase": "Phase 2", "description": "Implementation", "amount": 150000, "percent": "45%"},
            {"phase": "Phase 3", "description": "Completion", "amount": 83333, "percent": "25%"}
        ],
        "total_investment": 333333,
        "roi_projection": "Expected ROI of 180% within 18 months"
    }

def generate_excel_financial_enhanced(content: dict, structure: List[Section], company_name: str, job_id: str, request: ProposalRequest) -> str:
    """Generate enhanced Excel financial proposal"""
    try:
        filename = f"financial_proposal_{company_name.replace(' ', '_')}_{job_id}.xlsx"
        filepath = os.path.join(OUTPUT_DIR, filename)
        wb = openpyxl.Workbook()
        
        wb.remove(wb.active)
        
        ws_summary = wb.create_sheet("Executive Summary")
        create_summary_sheet(ws_summary, content, company_name)
        
        ws_payment = wb.create_sheet("Payment Schedule") 
        create_payment_schedule_sheet(ws_payment, content, company_name)
        
        ws_roi = wb.create_sheet("ROI Analysis")
        create_roi_analysis_sheet(ws_roi, content, company_name)
        
        wb.save(filepath)
        return filename
        
    except Exception as e:
        print(f"Error generating Excel financial document: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating Excel document: {str(e)}")

def create_summary_sheet(ws, content: dict, company_name: str):
    """Create executive summary sheet"""
    ws.merge_cells('A1:F1')
    ws['A1'] = f'Financial Proposal - {company_name}'
    ws['A1'].font = Font(bold=True, size=16, color="FFFFFF")
    ws['A1'].fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    
    ws.merge_cells('A3:F8')
    ws['A3'] = content.get('financial_summary', 'Financial summary not available')
    ws['A3'].alignment = Alignment(wrap_text=True, vertical='top')
    ws['A3'].font = Font(size=11)
    
    ws['A10'] = 'Total Investment:'
    ws['A10'].font = Font(bold=True)
    ws['B10'] = f"${content.get('total_investment', 0):,.2f}"
    
    ws['A11'] = 'ROI Projection:'
    ws['A11'].font = Font(bold=True)
    ws['B11'] = content.get('roi_projection', 'ROI details not available')
    
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width

def create_payment_schedule_sheet(ws, content: dict, company_name: str):
    """Create payment schedule sheet"""
    ws.merge_cells('A1:E1')
    ws['A1'] = f'Payment Schedule - {company_name}'
    ws['A1'].font = Font(bold=True, size=16, color="FFFFFF")
    ws['A1'].fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
    ws['A1'].alignment = Alignment(horizontal='center')
    
    headers = ['Phase', 'Description', 'Amount (USD)', 'Percentage', 'Due Date']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col)
        cell.value = header
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="336699", end_color="336699", fill_type="solid")
        cell.alignment = Alignment(horizontal='center')
    
    payment_schedule = content.get('payment_schedule', [])
    total_amount = 0
    
    for row_idx, payment in enumerate(payment_schedule, 4):
        ws.cell(row=row_idx, column=1, value=payment.get('phase', ''))
        ws.cell(row=row_idx, column=2, value=payment.get('description', ''))
        amount = payment.get('amount', 0)
        ws.cell(row=row_idx, column=3, value=amount)
        ws.cell(row=row_idx, column=4, value=payment.get('percent', ''))
        ws.cell(row=row_idx, column=5, value='Upon milestone completion')
        total_amount += amount
    
    total_row = len(payment_schedule) + 4
    ws.cell(row=total_row, column=1, value='TOTAL').font = Font(bold=True)
    ws.cell(row=total_row, column=3, value=total_amount).font = Font(bold=True)
    ws.cell(row=total_row, column=4, value='100%').font = Font(bold=True)

def create_roi_analysis_sheet(ws, content: dict, company_name: str):
    """Create ROI analysis sheet"""
    ws.merge_cells('A1:E1')
    ws['A1'] = f'ROI Analysis - {company_name}'
    ws['A1'].font = Font(bold=True, size=16, color="FFFFFF")
    ws['A1'].fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
    ws['A1'].alignment = Alignment(horizontal='center')
    
    ws['A3'] = 'ROI Projection:'
    ws['A3'].font = Font(bold=True, size=12)
    ws.merge_cells('B3:E3')
    ws['B3'] = content.get('roi_projection', 'ROI projection not available')
    
    headers = ['Year', 'Investment', 'Benefits', 'Net Benefit', 'Cumulative ROI']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=5, column=col)
        cell.value = header
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="336699", end_color="336699", fill_type="solid")
    
    roi_data = [
        ['Year 1', 200000, 150000, -50000, '-25%'],
        ['Year 2', 100000, 300000, 200000, '50%'],
        ['Year 3', 50000, 400000, 350000, '140%']
    ]
    
    for row_idx, data in enumerate(roi_data, 6):
        for col_idx, value in enumerate(data, 1):
            ws.cell(row=row_idx, column=col_idx, value=value)

@app.get("/status/{job_id}", response_model=ProposalResponse)
async def get_job_status_enhanced(job_id: str):
    """Get enhanced job status with structure information"""
    if job_id not in job_status:
        raise HTTPException(status_code=404, detail="Job not found")
    
    status_data = job_status[job_id]
    
    # Debug: Print files being returned
    files_list = status_data.get("files", [])
    if files_list:
        print(f"Status for {job_id}: returning {len(files_list)} files: {files_list}")
    
    response_data = {
        "job_id": job_id,
        "status": status_data["status"],
        "message": status_data["message"],
        "files": files_list,
        "progress": status_data.get("progress", 0)
    }
    
    if status_data["status"] != "error" and job_id in generated_structures:
        structure = generated_structures[job_id]
        response_data["structure_info"] = {
            "total_sections": len(ai_generator.flatten_sections(structure)),
            "main_sections": len(structure),
            "section_titles": [s.title for s in structure[:5]]
        }
    
    return ProposalResponse(**response_data)

@app.get("/structure/{job_id}")
async def get_generated_structure(job_id: str):
    """Get the generated structure for a specific job"""
    if job_id not in generated_structures:
        raise HTTPException(status_code=404, detail="Structure not found for this job")
    
    structure = generated_structures[job_id]
    
    return {
        "job_id": job_id,
        "structure": [section.to_dict() for section in structure],
        "total_sections": len(ai_generator.flatten_sections(structure)),
        "main_sections": len(structure)
    }

@app.get("/download/{filename}")
async def download_file(filename: str):
    """Download generated file"""
    file_path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type='application/octet-stream'
    )

@app.delete("/cleanup/{job_id}")
async def cleanup_job_enhanced(job_id: str):
    """Enhanced cleanup with structure removal"""
    if job_id in job_status:
        for filename in job_status[job_id].get("files", []):
            file_path = os.path.join(OUTPUT_DIR, filename)
            if os.path.exists(file_path):
                os.remove(file_path)
        
        del job_status[job_id]
        
        if job_id in generated_structures:
            del generated_structures[job_id]
        
        return {"message": "Job and associated data cleaned up successfully"}
    
    raise HTTPException(status_code=404, detail="Job not found")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "ai_model_available": ai_config.model is not None,
        "arabic_support": ARABIC_SUPPORT and AMIRI_FONT_PATH is not None,
        "active_jobs": len(job_status),
        "version": "2.1.0"
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)